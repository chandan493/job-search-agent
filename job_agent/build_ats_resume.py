#!/usr/bin/env python3
"""Build an ATS-friendly resume DOCX from the resume PDF and target job rows."""

from __future__ import annotations

import base64
import argparse
import hashlib
import json
import os
import re
import urllib.request
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Set, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
REPO_DIR = BASE_DIR.parent
CONFIG_PATH = Path(
    os.environ.get(
        "JOB_AGENT_CONFIG",
        str(REPO_DIR / "config.json" if (REPO_DIR / "config.json").exists() else BASE_DIR / "config.json"),
)
).expanduser()
DATA_DIR = CONFIG_PATH.parent / "data"
DEFAULT_OUTPUT_PATH = DATA_DIR / "tailored_resume.docx"

TAILORING_STOPWORDS = {
    "and", "are", "but", "can", "for", "from", "has", "have", "into", "our",
    "that", "the", "their", "this", "with", "you", "your", "role", "team",
    "work", "working", "experience", "years", "using", "used", "developer",
    "engineer", "software", "senior", "remote", "job", "jobs", "required",
    "requirements", "responsibilities", "preferred", "candidate", "company",
}
GENERIC_JOB_TERMS = {
    "software", "software engineer", "engineer", "developer", "senior software engineer",
}

SUPPORTED_SKILL_TERMS = [
    "Java", "J2EE", "Backend systems", "Spring Boot", "Spring MVC", "Spring Cloud", "Microservices",
    "REST APIs", "RESTful API", "API Management", "Kafka", "Event-driven systems",
    "Docker", "Kubernetes", "AWS", "Azure", "Azure Cosmos DB", "CI/CD", "Jenkins",
    "JUnit", "Mockito", "React", "React/Redux", "Angular", "AngularJS", "Couchbase",
    "DynamoDB", "RDS", "Lambda", "Camunda", "Rancher", "JFrog", "Grafana",
    "Performance optimization", "Application monitoring", "Code review",
    "Technical leadership", "Team mentoring", "Cloud-native development",
    "Distributed systems", "E-commerce", "Retail", "Healthcare insurance",
    "Telecom", "Order management", "Catalog services", "Generative AI",
    "GitHub Copilot", "Test automation",
]


def load_env_file(path: Path) -> None:
    if not path.exists():
        return
    for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        value = value.strip().strip('"').strip("'")
        if is_placeholder_secret(value):
            continue
        if key.strip() and key.strip() not in os.environ:
            os.environ[key.strip()] = value


def is_placeholder_secret(value: str) -> bool:
    lowered = value.strip().lower()
    return not lowered or "replace" in lowered or lowered in {"your-openai-api-key", "your-api-key"}


def post_json(url: str, payload: Dict[str, Any], api_key: str) -> Dict[str, Any]:
    request = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "User-Agent": "JobMatchAgent/1.0",
        },
        method="POST",
    )
    with urllib.request.urlopen(request, timeout=120) as response:
        return json.loads(response.read().decode("utf-8", errors="replace"))


def response_text(response: Dict[str, Any]) -> str:
    chunks: List[str] = []
    for item in response.get("output", []):
        for content in item.get("content", []):
            if content.get("type") == "output_text":
                chunks.append(content.get("text", ""))
    return "".join(chunks) or response.get("output_text", "")


def resume_cache_key(resume_path: Path, config: Dict[str, Any]) -> str:
    stat = resume_path.stat() if resume_path.exists() else None
    digest = ""
    if resume_path.exists():
        hasher = hashlib.sha256()
        with resume_path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                hasher.update(chunk)
        digest = hasher.hexdigest()
    raw = {
        "extract_version": 2,
        "path": str(resume_path.resolve() if resume_path.exists() else resume_path),
        "size": stat.st_size if stat else 0,
        "mtime_ns": stat.st_mtime_ns if stat else 0,
        "sha256": digest,
        "model": config.get("llm_resume_parser", {}).get("model", ""),
    }
    return hashlib.sha256(json.dumps(raw, sort_keys=True).encode("utf-8")).hexdigest()


def extract_resume(resume_path: Path, config: Dict[str, Any]) -> Dict[str, Any]:
    cache = DATA_DIR / "resume_full_ats_extract.json"
    cache_key = resume_cache_key(resume_path, config)
    if cache.exists():
        try:
            cached = json.loads(cache.read_text(encoding="utf-8"))
            if cached.get("cache_key") == cache_key and isinstance(cached.get("data"), dict):
                return cached["data"]
        except Exception:
            pass

    load_env_file(CONFIG_PATH.parent / ".env")
    load_env_file(BASE_DIR / ".env")
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key or is_placeholder_secret(api_key):
        raise RuntimeError(f"OPENAI_API_KEY is missing. Add it to {CONFIG_PATH.parent / '.env'}")

    model = config.get("llm_resume_parser", {}).get("model", "gpt-5.4-mini")
    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "name": {"type": "string"},
            "headline": {"type": "string"},
            "contact": {"type": "array", "items": {"type": "string"}},
            "summary": {"type": "string"},
            "skills": {"type": "array", "items": {"type": "string"}},
            "experience": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "title": {"type": "string"},
                        "company": {"type": "string"},
                        "location": {"type": "string"},
                        "dates": {"type": "string"},
                        "bullets": {"type": "array", "items": {"type": "string"}},
                    },
                    "required": ["title", "company", "location", "dates", "bullets"],
                },
            },
            "education": {"type": "array", "items": {"type": "string"}},
            "certifications": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["name", "headline", "contact", "summary", "skills", "experience", "education", "certifications"],
    }
    prompt = (
        "Extract the candidate's resume content faithfully from the attached PDF. "
        "Do not invent employers, dates, degrees, certifications, metrics, or contact details. "
        "Rewrite bullets only to be clearer and ATS-friendly while preserving the candidate's actual scope, domain, and seniority. "
        "Keep each bullet under 28 words. Preserve the candidate's actual facts."
    )
    file_data = base64.b64encode(resume_path.read_bytes()).decode("ascii")
    payload = {
        "model": model,
        "input": [
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": prompt},
                    {
                        "type": "input_file",
                        "filename": resume_path.name,
                        "file_data": f"data:application/pdf;base64,{file_data}",
                    },
                ],
            }
        ],
        "text": {
            "format": {
                "type": "json_schema",
                "name": "resume_full_extract",
                "strict": True,
                "schema": schema,
            }
        },
        "max_output_tokens": 5000,
    }
    data = json.loads(response_text(post_json("https://api.openai.com/v1/responses", payload, api_key)))
    cache.parent.mkdir(parents=True, exist_ok=True)
    cache.write_text(json.dumps({"cache_key": cache_key, "data": data}, indent=2, ensure_ascii=False), encoding="utf-8")
    return data


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def add_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(31, 78, 121)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.size = Pt(9.2)


def safe_filename(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("_")
    return value[:90] or "resume"


def candidate_slug(data: Dict[str, Any]) -> str:
    name = clean_phrase(data.get("name"))
    if not name:
        name = "candidate"
    return safe_filename(name).lower()


def tailored_resume_filename(data: Dict[str, Any], target_job: Optional[Dict[str, Any]] = None, job_id: Optional[str] = None) -> str:
    parts = [candidate_slug(data), "tailored_resume"]
    if target_job:
        company = safe_filename(clean_phrase(target_job.get("company", ""))).lower()
        title = safe_filename(clean_phrase(target_job.get("title", ""))).lower()
        if company:
            parts.append(company)
        if title:
            parts.append(title)
    if job_id:
        parts.append(safe_filename(job_id)[-18:].lower())
    return "_".join(part for part in parts if part) + ".docx"


def load_latest_job(job_id: str, latest_jobs_path: Path) -> Dict[str, Any]:
    payload = json.loads(latest_jobs_path.read_text(encoding="utf-8"))
    for job in payload.get("jobs", []):
        if job.get("job_uid") == job_id:
            return job
    raise KeyError(f"Job id not found in latest jobs: {job_id}")


def split_keywords(value: str) -> List[str]:
    return [item.strip() for item in re.split(r"[,;|]", value or "") if item.strip()]


def normalize_key(value: str) -> str:
    return re.sub(r"[^a-z0-9+#.]+", " ", value.lower()).strip()


def clean_phrase(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def unique_phrases(items: Sequence[Any]) -> List[str]:
    output: List[str] = []
    seen: Set[str] = set()
    for item in items:
        clean = clean_phrase(item)
        key = normalize_key(clean)
        if clean and key and key not in seen:
            seen.add(key)
            output.append(clean)
    return output


def resume_text_blob(data: Dict[str, Any]) -> str:
    parts: List[str] = [
        data.get("headline", ""),
        data.get("summary", ""),
        " ".join(map(str, data.get("skills", []) or [])),
    ]
    for item in data.get("experience", []) or []:
        parts.extend(
            [
                item.get("title", ""),
                item.get("company", ""),
                " ".join(map(str, item.get("bullets", []) or [])),
            ]
        )
    return " ".join(clean_phrase(part) for part in parts if part)


def job_text_blob(job: Dict[str, Any]) -> str:
    return " ".join(
        clean_phrase(job.get(key, ""))
        for key in ("title", "company", "location", "matched_keywords", "description")
        if job.get(key)
    )


def text_contains_phrase(text: str, phrase: str) -> bool:
    phrase_key = normalize_key(phrase)
    if not phrase_key:
        return False
    return phrase_key in normalize_key(text)


def supported_resume_terms(data: Dict[str, Any]) -> List[str]:
    resume_text = resume_text_blob(data)
    skills = data.get("skills", []) or []
    supported = [term for term in SUPPORTED_SKILL_TERMS if text_contains_phrase(resume_text, term)]
    supported.extend(skills)
    return unique_phrases(supported)


def job_requirement_terms(job: Dict[str, Any], data: Dict[str, Any]) -> List[str]:
    text = job_text_blob(job)
    matched = split_keywords(job.get("matched_keywords", ""))
    supported = supported_resume_terms(data)
    exact_hits = [term for term in supported if text_contains_phrase(text, term)]
    words = re.findall(r"[A-Za-z][A-Za-z0-9+#.]{2,}", text)
    word_hits = [
        word
        for word in words
        if word.lower() not in TAILORING_STOPWORDS and text_contains_phrase(resume_text_blob(data), word)
    ]
    return unique_phrases(matched + exact_hits + word_hits)[:45]


def canonical_supported_term(term: str, supported: Sequence[str]) -> str:
    key = normalize_key(term)
    for item in supported:
        item_key = normalize_key(item)
        if key == item_key or key in item_key or item_key in key:
            return item
    return clean_phrase(term)


def score_text(text: str, terms: Sequence[str], recent_weight: int = 0) -> int:
    lowered = normalize_key(text)
    score = recent_weight
    for index, term in enumerate(terms):
        key = normalize_key(term)
        if not key:
            continue
        if key in lowered:
            score += max(2, 10 - min(index, 8))
            if " " in key or "/" in key or "-" in key:
                score += 2
    if re.search(r"\b(\d+[%+]|\$\d+|\d+k\+|\d+\s*million|\d+\s*billion)\b", text, re.I):
        score += 3
    return score


def ranked_items(items: Sequence[str], terms: Sequence[str], limit: int) -> List[str]:
    scored: List[Tuple[int, int, str]] = []
    for idx, item in enumerate(items):
        scored.append((score_text(item, terms), -idx, item))
    ordered = [item for score, _neg_idx, item in sorted(scored, key=lambda row: (-row[0], row[1]))]
    return ordered[:limit]


def role_focus_title(base_headline: str, title: str, focus_terms: Sequence[str]) -> str:
    base = clean_phrase(base_headline) or "Experienced professional"
    focus = " | ".join(focus_terms[:4])
    if focus:
        return f"{base} | {focus} | Target: {title}"
    return f"{base} | Target: {title}"


def tailored_summary(
    original: str,
    headline: str,
    supported: Sequence[str],
    title: str,
    focus_terms: Sequence[str],
    role_fit: Sequence[str],
    limited_overlap: bool = False,
) -> str:
    focus = ", ".join(focus_terms[:8])
    proof = role_fit[0] if role_fit else ""
    base = clean_phrase(original)
    if not base:
        skill_text = ", ".join(unique_phrases(supported)[:8])
        base = clean_phrase(headline) or "Experienced professional"
        if skill_text:
            base = f"{base} with proven experience across {skill_text}."
    if focus and limited_overlap:
        base += f" This posting has limited direct keyword overlap, so this version emphasizes the closest truthful strengths: {focus}."
    elif focus:
        base += f" For this {title} role, the strongest alignment is {focus}."
    if proof:
        base += f" Relevant proof point: {proof}"
    return base if base else original


def tailor_experience(data: Dict[str, Any], terms: Sequence[str]) -> Tuple[List[Dict[str, Any]], List[str]]:
    tailored_experience: List[Dict[str, Any]] = []
    role_fit_pool: List[str] = []
    for exp_idx, item in enumerate(data.get("experience", []) or []):
        exp = dict(item)
        bullets = [clean_phrase(bullet) for bullet in item.get("bullets", []) or [] if clean_phrase(bullet)]
        recent_weight = max(0, 4 - exp_idx)
        scored = sorted(
            ((score_text(bullet, terms, recent_weight), -idx, bullet) for idx, bullet in enumerate(bullets)),
            key=lambda row: (-row[0], row[1]),
        )
        if exp_idx < 4:
            limit = 5
        elif exp_idx < 6:
            limit = 4
        else:
            limit = 3
        selected = [bullet for _score, _idx, bullet in scored[:limit]]
        exp["bullets"] = selected or bullets[:limit]
        tailored_experience.append(exp)
        role_fit_pool.extend(selected[:2])
    role_fit = ranked_items(role_fit_pool, terms, 4)
    return tailored_experience, role_fit


def tailor_for_job(data: Dict[str, Any], job: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not job:
        return dict(data)
    tailored = json.loads(json.dumps(data))
    title = job.get("title", "target role")
    terms = job_requirement_terms(job, tailored)
    non_generic_terms = [term for term in terms if normalize_key(term) not in GENERIC_JOB_TERMS]
    supported = supported_resume_terms(tailored)
    supported_keys = {normalize_key(item) for item in supported}
    job_focus_skills = [
        canonical_supported_term(term, supported)
        for term in non_generic_terms
        if any(normalize_key(term) == key or normalize_key(term) in key or key in normalize_key(term) for key in supported_keys)
    ]
    job_first_skills = list(job_focus_skills)
    job_first_skills.extend(term for term in supported if normalize_key(term) not in {normalize_key(item) for item in job_first_skills})
    limited_overlap = not bool(job_focus_skills)
    ranking_terms = non_generic_terms if non_generic_terms else unique_phrases(supported)[:8]
    tailored_experience, role_fit = tailor_experience(tailored, ranking_terms)
    focus_terms = unique_phrases(job_focus_skills)[:12] or unique_phrases(job_first_skills)[:6]
    tailored["skills"] = unique_phrases(job_first_skills)[:30] + [
        skill for skill in supported if normalize_key(skill) not in {normalize_key(item) for item in job_first_skills}
    ]
    tailored["experience"] = tailored_experience
    tailored["role_fit_bullets"] = role_fit
    tailored["headline"] = role_focus_title(tailored.get("headline", ""), title, focus_terms)
    tailored["summary"] = tailored_summary(tailored.get("summary", ""), tailored.get("headline", ""), supported, title, focus_terms, role_fit, limited_overlap)
    return tailored


def build_docx(data: Dict[str, Any], output_path: Path, target_job: Optional[Dict[str, Any]] = None) -> None:
    data = tailor_for_job(data, target_job)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.2)
    normal.paragraph_format.space_after = Pt(2)

    name = data.get("name") or "Candidate"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    r = title.add_run(name)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)

    headline = doc.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline.paragraph_format.space_after = Pt(1)
    run = headline.add_run(data.get("headline") or "Experienced professional")
    run.bold = True
    run.font.size = Pt(9.5)

    contact = " | ".join([c for c in data.get("contact", []) if c])
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        rr = p.add_run(contact)
        rr.font.size = Pt(8.5)
        rr.font.color.rgb = RGBColor(80, 80, 80)

    if target_job:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        rr = p.add_run(
            f"ATS-targeted for: {target_job.get('title', '')} at {target_job.get('company', '')}"
        )
        rr.italic = True
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor(90, 90, 90)

    add_section(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(data.get("summary", ""))
    run.font.size = Pt(9.2)

    if data.get("role_fit_bullets"):
        add_section(doc, "Target Role Fit")
        for bullet in data.get("role_fit_bullets", [])[:4]:
            add_bullet(doc, bullet)

    add_section(doc, "Core Skills")
    skills = data.get("skills", [])
    preferred = [
        "Java", "J2EE", "Spring Boot", "Spring MVC", "Spring Cloud", "Microservices", "REST APIs",
        "Kafka", "Docker", "Kubernetes", "AWS", "Azure", "CI/CD", "Jenkins", "JUnit", "Mockito",
        "React", "Angular", "Couchbase", "DynamoDB", "RDS", "Grafana", "API Management",
    ]
    merged = []
    skill_order = skills + preferred if target_job else preferred + skills
    for item in skill_order:
        clean = re.sub(r"\s+", " ", str(item)).strip()
        if clean and clean.lower() not in [x.lower() for x in merged]:
            merged.append(clean)
    table = doc.add_table(rows=0, cols=3)
    table.autofit = False
    remove_table_borders(table)
    for idx in range(0, min(len(merged), 30), 3):
        cells = table.add_row().cells
        for col in range(3):
            value = merged[idx + col] if idx + col < len(merged) else ""
            cells[col].width = Inches(2.42)
            set_cell_shading(cells[col], "F3F6FA")
            p = cells[col].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            rr = p.add_run(value)
            rr.font.size = Pt(8.7)

    add_section(doc, "Professional Experience")
    for job in data.get("experience", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        left = " - ".join([part for part in [job.get("title"), job.get("company")] if part])
        rr = p.add_run(left)
        rr.bold = True
        rr.font.size = Pt(9.4)
        meta = " | ".join([part for part in [job.get("location"), job.get("dates")] if part])
        if meta:
            rr2 = p.add_run(f"    {meta}")
            rr2.italic = True
            rr2.font.size = Pt(8.7)
            rr2.font.color.rgb = RGBColor(80, 80, 80)
        for bullet in job.get("bullets", [])[:5]:
            add_bullet(doc, bullet)

    if data.get("education"):
        add_section(doc, "Education")
        for item in data.get("education", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.add_run(item).font.size = Pt(9)

    if data.get("certifications"):
        add_section(doc, "Certifications")
        for item in data.get("certifications", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.add_run(item).font.size = Pt(9)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def build_resume_for_job(job_id: Optional[str] = None, output_path: Optional[Path] = None) -> Path:
    try:
        config = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise ValueError(
            f"Invalid JSON in {CONFIG_PATH}: line {exc.lineno}, column {exc.colno}. "
            "Check for a missing quote or comma, especially around resume_path."
        ) from exc
    resume_path = Path(config["resume_path"]).expanduser()
    if not resume_path.is_absolute():
        resume_path = CONFIG_PATH.parent / resume_path
    data = extract_resume(resume_path, config)
    target_job = None
    if job_id:
        target_job = load_latest_job(job_id, DATA_DIR / "latest_jobs.json")
    if output_path is None:
        if target_job:
            output_path = DATA_DIR / "generated_resumes" / tailored_resume_filename(data, target_job, job_id)
        else:
            output_path = DATA_DIR / tailored_resume_filename(data)
    build_docx(data, output_path, target_job)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Build an ATS-friendly resume DOCX.")
    parser.add_argument("--job-id", help="Generate a resume tailored to a job in data/latest_jobs.json")
    parser.add_argument("--output", help="Output DOCX path")
    args = parser.parse_args()
    output = build_resume_for_job(args.job_id, Path(args.output).expanduser() if args.output else None)
    print(str(output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
